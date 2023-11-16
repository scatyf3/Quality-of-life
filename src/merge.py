import os
import docx2pdf
from PyPDF2 import PdfFileMerger
from docx import Document
from datetime import datetime

# Convert lab_report.docx to lab_report.pdf
docx2pdf.convert("lab_report.docx")

# Modify cover.docx
doc = Document("cover.docx")

# Update line with custom string
custom_string = "Your Custom String" #在这里输入你的姓名学号
doc.paragraphs[0].text = custom_string

# Update line with system time
current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
doc.paragraphs[1].text = current_time

# Save modified cover.docx
doc.save("modified_cover.docx")

# Convert modified_cover.docx to modified_cover.pdf
docx2pdf.convert("modified_cover.docx")

# Merge modified_cover.pdf and lab_report.pdf into one PDF file
pdfs = ["modified_cover.pdf", "lab_report.pdf"]
merger = PdfFileMerger()

for pdf in pdfs:
    merger.append(pdf)

merger.write("merged.pdf")
merger.close( )

import os
import docx2pdf
from PyPDF2 import PdfFileMerger
from docx import Document
from datetime import datetime

# Convert lab_report.docx to lab_report.pdf
docx2pdf.convert("lab_report.docx")

# Modify cover.docx
doc = Document("cover.docx")

# Update line with custom string
custom_string = "Your Custom String" #在这里输入你的姓名学号
doc.paragraphs[0].text = custom_string

# Update line with system time
current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
doc.paragraphs[1].text = current_time

# Save modified cover.docx
doc.save("modified_cover.docx")

# Convert modified_cover.docx to modified_cover.pdf
docx2pdf.convert("modified_cover.docx")

# Merge modified_cover.pdf and lab_report.pdf into one PDF file
pdfs = ["modified_cover.pdf", "lab_report.pdf"]
merger = PdfFileMerger()

for pdf in pdfs:
    merger.append(pdf)

merger.write("merged.pdf")
merger.close( )

# Remove the intermediate files
os.remove("lab_report.pdf")

